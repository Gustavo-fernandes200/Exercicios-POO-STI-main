"""import os
import traceback
import pdfplumber
from docx import Document
from docx.shared import Inches
from PIL import Image
import pytesseract  # Biblioteca para OCR


def save_image(image_bytes, image_path):
    ""Salva os bytes da imagem em um formato compatível (PNG).""
    try:
        with Image.open(image_bytes) as img:
            img.save(image_path, format="PNG")
    except Exception as e:
        print(f"Erro ao salvar a imagem {image_path}: {e}")
        traceback.print_exc()


class PDFToWordConverter:
    ""Classe para converter PDF para Word.""

    def __init__(self, pdf_file_path, word_file_path, output_images_dir="images"):
        ""Construtor da classe.""
        self.pdf_file_path = os.path.normpath(pdf_file_path.strip('"'))
        if not os.path.isfile(self.pdf_file_path):
            raise FileNotFoundError(f"Arquivo {self.pdf_file_path} não encontrado")
        self.word_file_path = os.path.normpath(word_file_path.strip('"'))
        self.word_file_path = self.word_file_path if self.word_file_path.endswith(".docx") else self.word_file_path + ".docx"
        os.makedirs(os.path.dirname(self.word_file_path), exist_ok=True)
        self.output_images_dir = os.path.normpath(output_images_dir)
        if not os.path.isdir(self.output_images_dir):
            raise NotADirectoryError(f"O diretório {self.output_images_dir} não existe")

        os.makedirs(self.output_images_dir, exist_ok=True)  # Cria o diretório se não existir
        # Configura o caminho do executável do Tesseract (se necessário)
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

    def extract_images_from_pdf(self):
        ""Usa PDFplumber para extrair imagens do PDF.""
        try:
            with pdfplumber.open(self.pdf_file_path) as pdf:
                for page_number, page in enumerate(pdf.pages, start=1):
                    for image in page.images:
                        image_path = os.path.join(self.output_images_dir,
                                                  f"image_{page_number}_{image.page_number}.png")
                        save_image(image.image, image_path)
        except Exception as e:
            print(f"Erro ao processar PDF {self.pdf_file_path}: {e}")
            traceback.print_exc()

    def extract_text_from_image(self, image_path):
        ""Usa OCR para extrair texto de uma imagem.""
        try:
            with Image.open(image_path) as img:
                img = img.convert("RGB")
                img.save(image_path, format="PNG")
                return pytesseract.image_to_string(img, lang="por")  # Idioma configurado para português
        except Exception as e:
            print(f"Erro ao processar OCR na imagem {image_path}: {e}")
            traceback.print_exc()
            return ""

    def extract_text_from_pdf(self):
        ""Usa PDFplumber para extrair texto do PDF.""
        try:
            with pdfplumber.open(self.pdf_file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            print(f"Erro ao processar PDF {self.pdf_file_path}: {e}")
            traceback.print_exc()
            return ""

    def extract_tables_from_pdf(self):
        ""Usa PDFplumber para extrair tabelas do PDF.""
        try:
            with pdfplumber.open(self.pdf_file_path) as pdf:
                tables = []
                for page in pdf.pages:
                    tables.extend(page.extract_tables())
                return tables
        except Exception as e:
            print(f"Erro ao processar PDF {self.pdf_file_path}: {e}")
            traceback.print_exc()
            return []

    def extract_metadata_from_pdf(self):
        ""Usa PDFplumber para extrair metadados do PDF.""
        try:
            with pdfplumber.open(self.pdf_file_path) as pdf:
                metadata = pdf.metadata
                return metadata
        except Exception as e:
            print(f"Erro ao processar PDF {self.pdf_file_path}: {e}")
            traceback.print_exc()
            return {}

    def extract_data_from_pdf(self):
        ""Usa PDFplumber para extrair metadados e dados do PDF.""
        try:
            with pdfplumber.open(self.pdf_file_path) as pdf:
                metadata = pdf.metadata
                text = ""
                for page in pdf.pages:
                    text += page.extract_text()
                return metadata, text
        except Exception as e:
            print(f"Erro ao processar PDF {self.pdf_file_path}: {e}")
            traceback.print_exc()
            return {}, ""

    def extract_all_data_from_pdf(self):
        ""Usa PDFplumber para extrair metadados, dados e tabelas do PDF.""
        try:
            with pdfplumber.open(self.pdf_file_path) as pdf:
                metadata = pdf.metadata
                text = ""
                tables = []
                for page in pdf.pages:
                    text += page.extract_text()
                    tables.extend(page.extract_tables())
                return metadata, text, tables
        except Exception as e:
            print(f"Erro ao processar PDF {self.pdf_file_path}: {e}")
            traceback.print_exc()
            return {}, "", []

    def convert(self):
        # Cria um novo documento Word
        doc = Document()

        try:
            # Validação do arquivo PDF
            if not os.path.exists(self.pdf_file_path):
                raise FileNotFoundError(f"O arquivo PDF '{self.pdf_file_path}' não foi encontrado.")

            # Abre o arquivo PDF usando pdfplumber
            with pdfplumber.open(self.pdf_file_path) as pdf:
                if len(pdf.pages) == 0:
                    print("O PDF está vazio. Nenhuma página para converter.")
                    return

                for page_number, page in enumerate(pdf.pages, start=1):
                    doc.add_heading(f'Página {page_number}', level=1)

                    # Extrai imagem
                    image = page.images
                    if image:
                        image_path = os.path.join(self.output_images_dir, f"image_{page_number}.png")
                        save_image(image[0].image, image_path)
                        doc.add_picture(image_path)
                    else:
                        doc.add_paragraph("Imagem não disponível para esta página.")
                        print(f"Aviso: Não foi possível extrair imagem da página {page_number}.")

                    # Extrai texto
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
                    else:
                        doc.add_paragraph("Texto não disponível para esta página.")
                        print(f"Aviso: Não foi possível extrair texto da página {page_number}.")

                    # Extrai metadados
                    metadata = page.metadata
                    if metadata:
                        doc.add_heading(f'Metadados da Página {page_number}', level=2)
                        for key, value in metadata.items():
                            doc.add_paragraph(f"{key}: {value}")

                    else:
                        doc.add_paragraph("Metadados não disponíveis para esta página.")
                        print(f"Aviso: Não foi possível extrair metadados da página {page_number}.")

                    # Extrai tabelas
                    tables = page.extract_tables()
                    if tables:
                        doc.add_heading(f'Tabelas da Página {page_number}', level=2)
                        for table_index, table in enumerate(tables, start=1):
                            doc.add_paragraph(f'Tabela {table_index}:')
                            table_word = doc.add_table(rows=1, cols=len(table[0]))
                            table_word.style = 'Table Grid'
                            # Adiciona cabeçalho
                            for i, header in enumerate(table[0]):
                                table_word.cell(0, i).text = str(header)
                            # Adiciona linhas
                            for row in table[1:]:
                                row_cells = table_word.add_row().cells
                                for i, cell in enumerate(row):
                                    row_cells[i].text = str(cell)

                    # Extrai imagens
                    for image_index, image in enumerate(page.images, start=1):
                        image_bytes = image.get("stream").get_data()
                        image_path = os.path.join(
                            self.output_images_dir,
                            f"page_{page_number}_image_{image_index}.png"
                            if image_index == 1
                            else f"page_{page_number}_image_{image_index}.png",

                        )
                        # Salva a imagem no formato PNG
                        save_image(image_bytes, image_path)

                        # Adiciona imagem ao Word
                        doc.add_heading(f'Imagem da Página {page_number}, Nº {image_index}', level=2)
                        doc.add_picture(image_path, width=Inches(4.0))

                        # Extrai texto da imagem usando OCR
                        ocr_text = self.extract_text_from_image(image_path)
                        if ocr_text.strip():
                            doc.add_paragraph("Texto extraído da imagem:")
                            doc.add_paragraph(ocr_text)

            # Salva o documento Word
            doc.save(self.word_file_path)
            print(f"Arquivo convertido com sucesso: {self.word_file_path}")
            print(f"Imagens salvas no diretório: {self.output_images_dir}")

        except FileNotFoundError as e:
            print(f"Erro: {e}")
            traceback.print_exc()
        except PermissionError:
            print(f"Erro: Permissão negada para acessar o arquivo de saída '{self.word_file_path}'.")
            traceback.print_exc()
        except Exception as e:
            print(f"Ocorreu um erro inesperado: {e}")
            traceback.print_exc()


# Exemplo de uso
if __name__ == "__main__":

    pdf_file = input("Digite o caminho completo do arquivo PDF a ser convertido: ").strip()
    word_file = input(
        "Digite o caminho completo para salvar o arquivo Word (ex: C:\\Users\\ANTON\\Documents\\resultado.docx): ").strip()

    # Adiciona um nome padrão ao arquivo Word caso o usuário não especifique um nome
    if not word_file.endswith(".docx"):
        word_file = os.path.join(word_file, "resultado.docx")
        print(f"Não foi especificado um nome de arquivo. Salvando como: {word_file}")

    converter = PDFToWordConverter(pdf_file, word_file)
    converter.convert()
"""

import os
import traceback
import pdfplumber
from docx import Document
from docx.shared import Inches
from PIL import Image
import pytesseract  # Biblioteca para OCR
import io


def save_image(image_bytes, image_path):
    """Salva os bytes da imagem em um formato compatível (PNG)."""
    try:
        # Converte os bytes da imagem para uma imagem PIL
        with Image.open(io.BytesIO(image_bytes)) as img:
            img.save(image_path, format="PNG")
            print(f"Imagem salva com sucesso: {image_path}")
    except Exception as e:
        print(f"Erro ao salvar a imagem {image_path}: {e}")
        traceback.print_exc()


class PDFToWordConverter:
    """Classe para converter PDF para Word."""

    def __init__(self, pdf_file_path, word_file_path, output_images_dir="images"):
        """Construtor da classe."""
        self.pdf_file_path = os.path.normpath(pdf_file_path.strip('"'))
        if not os.path.isfile(self.pdf_file_path):
            raise FileNotFoundError(f"Arquivo {self.pdf_file_path} não encontrado.")

        # Configura caminhos de saída
        self.word_file_path = os.path.normpath(word_file_path.strip('"'))
        if not self.word_file_path.endswith(".docx"):
            self.word_file_path += ".docx"
        os.makedirs(os.path.dirname(self.word_file_path), exist_ok=True)

        self.output_images_dir = os.path.normpath(output_images_dir)
        os.makedirs(self.output_images_dir, exist_ok=True)

        # Configura o caminho do executável do Tesseract (se necessário)
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

    def extract_text_from_image(self, image_path):
        """Usa OCR para extrair texto de uma imagem."""
        try:
            with Image.open(image_path) as img:
                return pytesseract.image_to_string(img, lang="por")  # Idioma configurado para português
        except Exception as e:
            print(f"Erro ao processar OCR na imagem {image_path}: {e}")
            traceback.print_exc()
            return ""

    def convert(self):
        """Realiza a conversão do PDF para Word."""
        doc = Document()

        try:
            with pdfplumber.open(self.pdf_file_path) as pdf:
                if len(pdf.pages) == 0:
                    print("O PDF está vazio. Nenhuma página para converter.")
                    return

                for page_number, page in enumerate(pdf.pages, start=1):
                    # Adiciona cabeçalho para a página
                    doc.add_heading(f'Página {page_number}', level=1)

                    # Extrai texto da página
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
                    else:
                        doc.add_paragraph("Texto não disponível para esta página.")
                        print(f"Aviso: Não foi possível extrair texto da página {page_number}.")

                    # Extrai tabelas da página
                    tables = page.extract_tables()
                    if tables:
                        doc.add_heading(f'Tabelas da Página {page_number}', level=2)
                        for table_index, table in enumerate(tables, start=1):
                            doc.add_paragraph(f'Tabela {table_index}:')
                            table_word = doc.add_table(rows=1, cols=len(table[0]))
                            table_word.style = 'Table Grid'

                            # Adiciona cabeçalhos e linhas à tabela
                            for i, header in enumerate(table[0]):
                                table_word.cell(0, i).text = str(header)
                            for row in table[1:]:
                                row_cells = table_word.add_row().cells
                                for i, cell in enumerate(row):
                                    row_cells[i].text = str(cell)

                    # Extrai imagens da página
                    for image_index, image in enumerate(page.images, start=1):
                        try:
                            image_bytes = image.get("stream").get_data()
                            image_path = os.path.join(
                                self.output_images_dir, f"page_{page_number}_image_{image_index}.png"
                            )
                            save_image(image_bytes, image_path)

                            # Adiciona a imagem ao documento Word
                            doc.add_heading(f'Imagem da Página {page_number}, Nº {image_index}', level=2)
                            doc.add_picture(image_path, width=Inches(4.0))

                            # Extrai texto da imagem usando OCR
                            ocr_text = self.extract_text_from_image(image_path)
                            if ocr_text.strip():
                                doc.add_paragraph("Texto extraído da imagem:")
                                doc.add_paragraph(ocr_text)

                        except Exception as e:
                            print(f"Erro ao processar imagem na página {page_number}: {e}")
                            traceback.print_exc()

            # Salva o documento Word
            doc.save(self.word_file_path)
            print(f"Arquivo convertido com sucesso: {self.word_file_path}")
            print(f"Imagens salvas no diretório: {self.output_images_dir}")

        except FileNotFoundError as e:
            print(f"Erro: Arquivo PDF '{self.pdf_file_path}' não encontrado.")
            traceback.print_exc()
        except PermissionError as e:
            print(f"Erro: Permissão negada para acessar o arquivo de saída '{self.word_file_path}'.")
            traceback.print_exc()
        except Exception as e:
            print(f"Ocorreu um erro inesperado: {e}")
            traceback.print_exc()


# Exemplo de uso
if __name__ == "__main__":
    pdf_file = input("Digite o caminho completo do arquivo PDF a ser convertido: ").strip()
    word_file = input(
        "Digite o caminho completo para salvar o arquivo Word (ex: C:\\Users\\ANTON\\Documents\\resultado.docx): "
    ).strip()

    # Configura nome padrão do arquivo Word, se necessário
    if not word_file.endswith(".docx"):
        word_file = os.path.join(word_file, "resultado.docx")
        print(f"Não foi especificado um nome de arquivo. Salvando como: {word_file}")

    converter = PDFToWordConverter(pdf_file, word_file)
    converter.convert()

